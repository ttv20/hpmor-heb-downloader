import requests
import collections
import os
import docx

##############
# parameters #
##############

# the name of the file containing the entire book, which the program will create
unified_file_name = 'hpmor-heb'

# the file format for all downloaded files, can also be doc but requires tinkering
file_format = '.docx'

# the name of the folder which will contain all the downloaded single chapters
new_folder_name = 'single-chapters'

# a list of tuples, each containing the number of a chapter and its google drive id
chapter_ids = [('001', '1Y5XeiiJQMBtjiRGJSFZL__LelONoR-uONcsRYenzb_k'),
               ('002', '1IT5bEcnhO4cF4zw36Bbrcvq71x88UN-F1k3rYbtG4d8'),
               ('003', '18-iuESSqMYSd_6T3FFOn-Vgn2BzgNWBNDtCG_pI-H78'),
               ('004', '1o90K4NgkU-15hPSZnyu0ZyCKAyxytaqBz3B4y6jpJi8'),
               ('005', '1RvV9wVmTR6k5THwdXICDWsiX_cLy9gkegkqgQzTr8WU'),
               ('006', '1l_qBTC6CEMXFErXvdADM8UuMQH0VuW2EV4e-r2QRdI4'),
               ('007', '1muj9kQY54P-Fip5IXsjK-_H_yv-bI5b377rpvxhhF08'),
               ('008', '1LvZTbzv9ASoyiNk7py3vMgTUnm5i8dm8aOzU84Z_6tg'),
               ('009', '1yCvLSAZTqrokYmeC6NLXgZW_I98LQPKEmj9kQvqwG_g'),
               ('010', '1vlC7MUC9e6zB9dW1k6clDZhrjquxamo7s7ThgR03USo'),
               ('011', '1w7YFgmeXIlgv6NhadDCo9ikAZoxsdo7ByN-gBe9-WHU'),
               ('012', '1G7Xd9ed66s47PrBcdiFGB0I-VFUn5flG0KUSsvpSElY'),
               ('013', '1f4TfoTA4B4sTR_SaZf7kw-WO-5gCGPDmXka3WtIJ8gY'),
               ('014', '13vhWPYEtHhmEYqoPJFfyMkwOqpo0VaXxHhRpcSMD00o'),
               ('015', '1JVLSLJEjiwYbCegX_HtuDsKVZWPYfWv1VmlwNckcdEs'),
               ('016', '1m_LTKdWphm7IP1D73qGNjp9eMo1zqAWcirQGmgryg-w'),
               ('017', '1sY-DaSIuWf29FlsyQnmtj_oHRPqUjRdujSnFLuxdWsU'),
               ('018', '1PrNoScMvBcqqHp6mgp9_6yS2uA8_4CX3lNZ94J93yZc'),
               ('019', '1W0SKPiXH5msYtZRU3y6MAMeYbpMiq_TtDJdADLhqMAA'),
               ('020', '1xsredct16S5xHcnYM-rZq0Z2WEBvbplOsjC4F26XeDU'),
               ('021', '1q5OS8LoxKM9OZk85T8wyi5gNnBS3KBkTgG6F4TifASA'),
               ('022', '1hxskHtBlFhIWXw4d9GT4757z7twsaOrgV41Wen35G9o'),
               ('023', '1X9WNBJDFjaWDICGHb0yjEzzJdtyZ1dqU4J1IduPr3Ro'),
               ('024', '1Auq6iQJI96gk3Ico8kNVbYdFljX1lCwMhohEPEdA6po'),
               ('025', '12NtTpJ1pwDhGhekAiyqGGCZsdd8Fpl9tXmiXecXkrdY'),
               ('026', '16LD731SYMEnlm8mIr7fragb9WvmNnqcIM0tnJtbNIMY'),
               ('027', '1WGX9WfXBvZ9W3KWAHh_-6_lmgMILTmmRnDA8SM-lZRo'),
               ('028', '1jwnkLTcdBfLxoDg0fOWGqZmBDFo6RIeYmOajFPJnPcI'),
               ('029', '1vobFENvFaweaMDpRVWOrHi59AnkW72aDw7zBsCi6-8E'),
               ('030', '1ODZZ9d2EtP12Bckh21GwESdIN05vhKgwgTQz7woPdP4'),
               ('031', '1My_yL2uo5Aqf4YDdlbZg6d-IELygunNo2vHYvwueR2E'),
               ('032', '1Z0bNKRBkRK-ZkPK5dYPXWjuvft28WVAXy1cqXXOtX2k'),
               ('033', '1c7suHN3GVn8ZP4Qx_cK70Ee1evxth1IxeT308KAxj1Q'),
               ('034', '1CW8hRSK0BQNWbELweqJ9JnmbxdmEMSYn81iAVbegAfY'),
               ('035', '13CHZ5th3mOkxlNeP-kuXbTFR0XyOqFUINVsH4UuemIg'),
               ('036', '1HDMxmDbHP8a9z5-UdpmrwVSycUTRFDKV7kPYNZJPZCI'),
               ('037', '1fKyuXHW1jgO4T5kekasGykCwFFU9dgDCQmJPKm-6-VU'),
               ('038', '1MouCmQvJfuvTGIC1pEIf1THOygyAWrYMgeDwHC4f5kA'),
               ('039', '18LMs0J4M1W8FrRINhm7aLAVvqjy_SDdP6M38SkMCNdg'),
               ('040', '1LK4YrR2kZrLUQY0dyxCAvQzT39M6KdCRZnlUhmUL5mE'),
               ('041', '16R5tc4OJCU82HWtcX_9yQCTkBOl2Qum4cHvWSI7dGYg'),
               ('042', '1MDZHzuIf5NTpPjKU3xTM_wTXxW2Yj7HkSqD_OdZ_1lk'),
               ('043', '1BRwuhwElqmZD0ey7wGY1HyTD1H1hrOW9h_MIa3LpmIM'),
               ('044', '1P2yG98xdG2rt9YR2iMLpE8_JgtcVR_tEhB8lPJ1Pwaw'),
               ('045', '1aGqiRYjFoGDY_1sFpQfhUrFBBEONFZd48ydfnntJ1zA'),
               ('046', '1yQqqJvzclAjyUXfuYnUwNGa4n3KhwwkzpjEnnQYKGU8'),
               ('047', '16t3yIHAQYK6CneShv3agrm9S-ICY4e-mMGK77Y7kL8I'),
               ('048', '1eYCiWztswwucrZdQWTa4QmD-2gHYSzCkAELgO4P5T1I'),
               ('049', '1JaKt3gSK0KrJNRwu40HQI-GF7EAqKySWqJ-LPglv5Mk'),
               ('050', '19TNqg6g_dlefrmrG8IjooJKvcLov7tI0vRXQ1zbZr7A'),
               ('051', '1tR951wQtbqnDbHMfXGYC45uP6DsRuclAVrbyaV6V_08'),
               ('052', '192FlydC_-tNEAFmT4Nf7mGMmLObiBJb1Pnvk9zgADyE'),
               ('053', '1FwFCAmcxd5UR-2XsYUjqIiBXlaZIZGKr57EILKbJpJs'),
               ('054', '1lr6RP1B_YHM3oYDtmdRQe_TDFlkX4V_vAru1ETWNb4k'),
               ('055', '1-qITWvbdXbrPIvRmB4OqiAaoUiSMoY-OH2r7jvlaSKY'),
               ('056', '10GtFLJYVbWmHcdyOgqTIH5pzNs2W5wHllLP0-zU7utY'),
               ('057', '1Ryug8CrRI5NBcN3WRnbknHortxuoefpU-NQ9o2arEbQ'),
               ('058', '1R3CpU-DU_KK68b1LKv8S21wY6wo586uVahBw5UwKMF0'),
               ('059', '1kebX06pq6ET2G-NVzxpv3Ot2Eb6qxNu_9-zw-arLY_o'),
               ('060', '12zT8Zmehii98rfN1u3GBjCyhH09oKmOEqQ__xYvgpP0'),
               ('061', '1y2Siqj4F6t-YJXP-nP_pXDc134rfdslKJ_G4JukYNdM'),
               ('062', '1m-3PDYN7zOAYHZfAIZ57jpTv0jZ9ALgtMfFVn8U8Sa4'),
               ('063', '1B-ZG0Fi1mBhgnQMeZAu1N5vWaLYrl8W_so8yoM8dayk'),
               ('064', '1vo-S9nzFVvD65-PAS8CNLEyuh3Z9gPNdm5hqArV-Dns'),
               ('065', '1c4YOkNovyCqehB1XHixhounfsF3tezmnBv8AG61sHD8'),
               ('066', '1lyv-6pYr9_P1dhz9OBm3R06gl1f_dIkoD1eSa2i0LJU'),
               ('067', '1gqYG5JmdddJqHVwL7i8IWMOqVuGWiKVsZpIVe9409vg'),
               ('068', '1CgpMdrVznJecut7hnyg-63RTzAYCAa-cX4CCnVoiU8g'),
               ('069', '1mVaaxPrDp6wMJccJDLBiAe-gUYJ8UIb5PbNCKj3UxOA'),
               ('070', '1cDdyZEID68nzFMa1TBuVA7tDdBQA_o6HFPTS9UrxI-4'),
               ('071', '1Ek1Ez25M_nQfPLSeMQc0B97UePK-KZixkpKvoPCYZTE'),
               ('072', '1CW0cG0aitO8caP2PBEdsoIdW9AeLIGJZJi_BALCbSeg'),
               ('073', '1M6nJds0ZLUw4a8mriB0Y2xlu8qsiRa_67ee0iYHKJnI'),
               ('074', '16ErgE98VdYm-P9jXalzguwMZZm9XWXVzbaE1nB43a3Y'),
               ('075', '133vopqHqsXAr7_Q9ahTw31wmdDjuFASF311OQoQtiWc'),
               ('076', '1KBxk775a9tDimYk8gWNFvMfCScQ14bTZvw_Rn14fA_M'),
               ('077', '1s091Z6eYYg1cesZAVeqfGFXwoTGVLYsNn-bx3jCPskM'),
               ('078', '10PcMYH_DHcT7UyWkUKVUaUNAWeUYYXjnyzY_zLtfoLc'),
               ('079', '1_gRfp5-88DnN-oxXIWEOM-J7jGkyKHoWlVGlgMTwF0w'),
               ('080', '1cGC4gEJYJpWxHjphY3HDYhqBnB5TX6EzbgrFsN1WulY'),
               ('081', '13f1A8nqcw2JmqKb27CldS1iIoVzdTxotpYtsBNb2l-Y'),
               ('082', '1bzgv9tCCzhR8WV9S7Gzjze-UuB20T-9BIPoy_luQmvA'),
               ('083', '1nd20rDQ-rPid1mPqurXsRvrWKz32objYXgrdanA19AA'),
               ('084', '146YdEXW9fhYsYY-E3bOpAtCALYbgtt3HXnCdQDXIApM'),
               ('085', '1c0dUVi-D_VovEMzAj1BgRheo9n0eKUb5AMW0lTnGqd4'),
               ('086', '1cpaZIGyHq-QrbrbCuysqpVHaeN7EMfGFBxsmYS1T5mU'),
               ('087', '13I2Lb0p5lNpJUOF4iz8zmf8IV0k-xzOXqiCzhgLcosE'),
               ('088', '1XPvwbdItioTADuRVJSAHwvcxuAXH3Fhea22aDGVUDjM'),
               ('089', '17oT8g34OqeYzyjS-c6MlTvTUm4fMcorWBjxKllYbE5Q'),
               ('090', '1sEiEfsr1APan6yodJcWvrIN5lzA1OCuZHr6noLWB0QE'),
               ('091', '10OuExkgmXpbnKwb0hSDhbG4vmwYhB9CUcwdfY0FPQQo'),
               ('092', '1aORaK5OMbfuuBG6EMIKkByzFRWHVec6uD1hl630LHbU'),
               ('093', '17pJ-i-fkuAqJPmi82WYaK9ih-yFtszCinRwBrEzFpXI'),
               ('094', '159DtG7ikQ5DOwnT9-3LIzdkg0XEW2mWoksM2YNmz7iU'),
               ('095', '1j1l9qsQJ3xLuSTbvuLAXnw6QO2OA4qx1MOuPF9PWp9w'),
               ('096', '1yjd4oYxzMzxqwMJb2h9nlVBK40Bln_q9Cq_QuIH8Exs'),
               ('097', '1ZhNUCGJii5OxeW6_TrHu9ltZc1JXxAbt3cgZy23wOQY'),
               ('098', '1OOzzwRKMHdOsTrx6TQPKgJ74KsICHSAuvLvcwCZADfs'),
               ('099', '1jzicVTGSR0G6PHGvxDadSjSaoylkMsX-ma3MUHAjlDI'),
               ('100', '1qaXETgnzUc5i9EIqGMbKBRjGafSZZcGy1JYQ8wIpp5E'),
               ('101', '1fh1KIGwHWe2sH78WuY6ywc-NdYoJ7vzHu4o1er8T680'),
               ('102', '19ApD1q5bfQrTKzJSENC4rxexGDL1EzW7Eg-QjlxbBdM'),
               ('103', '12dq8rGdi_bP8X8kOUMOwB7wPsybXn1WVWLGq7tmJmzY'),
               ('104', '1MBqHRVlZ6hHim3FXmU_1O9DfwnmtPi8mepwv4eMEmQU'),
               ('105', '1gF5Gm9qsr3aIkDD33LtTojKIpvNXm9JxKC-VjgO8rYs'),
               ('106', '1bKEap05PXFrzy32ygMHtbKitkRYhePam39Jo-Q9xOWI'),
               ('107', '1a4yVCBOrdzTvWZMjmRcgUzdbaFrmEoFfcx18yLEJM-A'),
               ('108', '1oZxuReWkirHHsDAfFeSXhEyPv-NokhyFEM_OwTajBzA'),
               ('109', '19NC0vv1Bf7AggwBT1-vRi_uBIbg6NtqalmXhuOcoSm8'),
               ('110', '117Rtc2H8OJxs_NdQnLF22wf2BzXo4pDbKdcKGrTrBW0'),
               ('111', '1YLf17jIUZuw6aCpUt6UyBY-zAfkX_WYEkXz4y3UHX3k'),
               ('112', '1RnunthktZ5-E8eh7cDh_xT0ybqRu-VPMCgaJ2nzzr84'),
               ('113', '1widaXdtjeRtGbkrghp6q2f-VncSPV75kTkacRohiYS8'),
               ('114', '1Lc0Kcn5Dc9tOpXAv7vBvivWErJcyCulVeEdGNjm888g'),
               ('115', '1kJiHCCoMEFM3d3t_GPiAw5URW-mafBLVI_v30WZdbT0'),
               ('116', '1I0J43AZvC5qrVqxNo74tTUy0WvADd52kFEbFEMQy7No'),
               ('117', '1Stf8bExF7HYyRwoQgonLIeBYA1AVjRWA6Cw6tWS1_qM'),
               ('118', '1E2W1hWGkEJ9TkyTEjqMssXtXw1Mpz8OgW4XSWm63-eY'),
               ('119', '1i4hkqgcC-U6yofQlCaIFA0IK4mtejidwiKdJDtqSxdc'),
               ('120', '1s7hwn5_Z6E-3Ror-oavafjtjxbcpkCzPuCrdXBmWHRw'),
               ('121', '1vk9DhmY7uPuVQ5UayatrAEkpIoUWiJl-FTbGle_Js8c'),
               ('122', '1dX1ADASKwJXU2RBrDeu9cCdp0gEuofrTrfjbbca20EQ')]

# an ordered dictionary made from the list of tuples above, keys: chapter numbers, values: chapter google ids
chapters_dict = collections.OrderedDict(chapter_ids)


#############
# functions #
#############

# downloads a single google docs document using its google id
def single_downloader(file_id, file_name, destination):

    # getting the correct url
    default_url = 'https://docs.google.com/document/d/FILE_ID/export?format=doc'
    url = default_url.replace('FILE_ID', file_id)

    # downloading the file
    r = requests.get(url, allow_redirects=True)

    # saving the file
    open(destination + '\\' + file_name + file_format, 'wb').write(r.content)
    print('chapter ' + str(file_name) + ' downloaded')


# downloads all the hpmor chapters using the chapter dictionary and single_downloader
def complete_downloader(chapter_dict, destination):

    # a list of chapter file names for later use
    files = []

    for key, value in chapter_dict.items():
        single_downloader(value, key, destination)
        files.append(key + file_format)

    print('\nall chapters downloaded successfully\n')
    return files


# a version of complete_downloader with the paths automatically set to the current folder in which the py is in
def complete_downloader_pathless(chapter_dict):

    # getting the absolute path of the current folder and adding the new folder to create
    new_path = os.path.dirname(os.path.abspath(__file__))
    new_path = new_path + '\\' + new_folder_name

    if not os.path.exists(new_path):
        os.makedirs(new_path)
        print('created folder: ' + new_folder_name + '\n')
    else:
        print('There is already a folder named ' + new_folder_name + ' in this directory.\n'
                                                                     'Please delete it and try again.')

    files = []

    for key, value in chapter_dict.items():
        single_downloader(value, key, new_path)
        files.append(new_path + '\\' + key + file_format)

    print('\nall chapters downloaded successfully\n')
    return files


# merges all downloaded chapters to a single word document
def complete_unifier(chapter_dict, destination):
    doc_list = complete_downloader(chapter_dict, destination)
    doc_list.sort()

    # the first chapter is the one all others are added to in the following loop
    first_doc = docx.Document(doc_list[0])
    first_doc.add_page_break()

    # iterating over all chapter files except the first one
    for other in doc_list[1:]:

        # opening the other chapter and appending a page break to it
        other_doc = docx.Document(other)
        other_doc.add_page_break()

        # adding the contents of the other chapter to the first
        for element in other_doc.element.body:
            first_doc.element.body.append(element)

        print('chapter ' + other + ' merged')

    # saving the newly created merged document
    first_doc.save(unified_file_name + file_format)

    print('\nall files merged successfully')


# a version of complete_unifier with the paths automatically set to the current folder in which the py is in
def complete_unifier_pathless(chapter_dict):

    # only difference - using the pathless version of complete_downloader
    doc_list = complete_downloader_pathless(chapter_dict)
    doc_list.sort()

    first_doc = docx.Document(doc_list[0])
    first_doc.add_page_break()

    for other in doc_list[1:]:
        other_doc = docx.Document(other)
        other_doc.add_page_break()

        for element in other_doc.element.body:
            first_doc.element.body.append(element)

        print('chapter ' + other + ' merged')

    first_doc.save(unified_file_name + file_format)

    print('\nall files merged successfully')


########
# main #
########

complete_unifier_pathless(chapters_dict)
